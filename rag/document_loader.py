"""
document_loader.py — Loads documents from various file formats into LangChain Documents.
"""
from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import List

from langchain_core.documents import Document

from campusmind.config.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """
    Adapter that converts uploaded files into LangChain Document objects.
    Supports: PDF, TXT, DOCX, CSV, MD.
    """

    @staticmethod
    def load_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
        """
        Load a document from raw bytes and filename.

        Args:
            file_bytes: Raw file content.
            filename: Original filename (used to detect format).

        Returns:
            List of LangChain Document objects.
        """
        ext = Path(filename).suffix.lower().lstrip(".")
        logger.info("Loading document: %s (type=%s)", filename, ext)

        loaders = {
            "pdf": DocumentLoader._load_pdf,
            "txt": DocumentLoader._load_txt,
            "md": DocumentLoader._load_txt,
            "docx": DocumentLoader._load_docx,
            "csv": DocumentLoader._load_csv,
        }

        loader_fn = loaders.get(ext)
        if loader_fn is None:
            raise ValueError(f"Unsupported file type: .{ext}")

        docs = loader_fn(file_bytes, filename)
        logger.info("Loaded %d chunks from %s", len(docs), filename)
        return docs

    @staticmethod
    def _load_pdf(data: bytes, filename: str) -> List[Document]:
        """Extract text from PDF pages using pypdf."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            docs = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": filename, "page": i + 1},
                    ))
            return docs
        except Exception as exc:
            logger.error("PDF load error: %s", exc)
            raise

    @staticmethod
    def _load_txt(data: bytes, filename: str) -> List[Document]:
        """Load plain text or markdown file."""
        text = data.decode("utf-8", errors="replace")
        return [Document(page_content=text, metadata={"source": filename})]

    @staticmethod
    def _load_docx(data: bytes, filename: str) -> List[Document]:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return [Document(page_content=text, metadata={"source": filename})]
        except Exception as exc:
            logger.error("DOCX load error: %s", exc)
            raise

    @staticmethod
    def _load_csv(data: bytes, filename: str) -> List[Document]:
        """Convert CSV rows into Document objects."""
        text_io = io.StringIO(data.decode("utf-8", errors="replace"))
        reader = csv.DictReader(text_io)
        docs = []
        for i, row in enumerate(reader):
            content = "\n".join(f"{k}: {v}" for k, v in row.items())
            docs.append(Document(
                page_content=content,
                metadata={"source": filename, "row": i + 1},
            ))
        return docs
