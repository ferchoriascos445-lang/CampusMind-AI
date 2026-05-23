"""
report_tool.py — Generates PDF, DOCX, TXT, and Markdown export files.
Soporta renderizado de ecuaciones LaTeX en PDF y Word.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path

from config.logger import get_logger
from config.settings import settings

logger = get_logger(__name__)


def _split_latex(text: str) -> list[dict]:
    """
    Divide el texto en bloques tipados:
    - latex_block: ecuación $$ ... $$ o \\[ ... \\]
    - latex_inline: ecuación $ ... $ o \\( ... \\)
    - heading/subheading/title/bullet/numbered/empty/paragraph
    """
    blocks = []

    # Primero separar por bloques LaTeX de bloque ($$...$$)
    parts = re.split(r'(\$\$[\s\S]*?\$\$|\\\[[\s\S]*?\\\])', text)

    for part in parts:
        if not part:
            continue

        # Bloque LaTeX
        if (part.startswith('$$') and part.endswith('$$')) or \
           (part.startswith('\\[') and part.endswith('\\]')):
            formula = part.strip('$').strip()
            if part.startswith('\\['):
                formula = part[2:-2].strip()
            blocks.append({"type": "latex_block", "text": formula})
            continue

        # Procesar línea a línea el texto normal
        for line in part.split("\n"):
            raw = line.rstrip()
            if not raw.strip():
                blocks.append({"type": "empty"})
            elif raw.startswith("### "):
                blocks.append({"type": "subheading", "text": raw[4:].strip()})
            elif raw.startswith("## "):
                blocks.append({"type": "heading", "text": raw[3:].strip()})
            elif raw.startswith("# "):
                blocks.append({"type": "title", "text": raw[2:].strip()})
            elif re.match(r"^[-*•]\s+", raw):
                blocks.append({"type": "bullet", "text": re.sub(r"^[-*•]\s+", "", raw)})
            elif re.match(r"^\d+\.\s+", raw):
                blocks.append({"type": "numbered", "text": re.sub(r"^\d+\.\s+", "", raw)})
            elif raw.startswith("**") and raw.endswith("**") and len(raw) > 4:
                blocks.append({"type": "heading", "text": raw.strip("*").strip()})
            else:
                blocks.append({"type": "paragraph", "text": raw})

    return blocks


def _latex_to_image(formula: str, fontsize: int = 14, dpi: int = 150) -> bytes | None:
    """
    Convierte una fórmula LaTeX a imagen PNG usando matplotlib.
    Retorna bytes de la imagen o None si falla.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.mathtext as mathtext

        fig, ax = plt.subplots(figsize=(8, 1.2))
        ax.axis("off")

        # Limpiar la fórmula para matplotlib
        clean = formula.strip().strip('$')

        ax.text(
            0.5, 0.5,
            f"${clean}$",
            ha="center", va="center",
            fontsize=fontsize,
            transform=ax.transAxes,
        )

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as exc:
        logger.warning("LaTeX→image failed: %s", exc)
        return None


class ReportTool:
    def __init__(self) -> None:
        output_dir = Path(settings.OUTPUTS_DIR)
        output_dir.mkdir(parents=True, exist_ok=True)
        self._output_dir = output_dir

    def export_chat(self, messages: list[dict], fmt: str = "txt", title: str = "Chat Export") -> bytes:
        text_lines = []
        for m in messages:
            role = "You" if m["role"] == "user" else "CampusMind AI"
            text_lines.append(f"{role}:\n{m['content']}\n")
        full_text = "\n".join(text_lines)
        return self.generate_report(full_text, title, fmt)

    def generate_report(self, content: str, title: str, fmt: str = "pdf") -> bytes:
        dispatch = {
            "pdf": self._to_pdf,
            "docx": self._to_docx,
            "txt": self._to_txt,
            "md": self._to_md,
        }
        fn = dispatch.get(fmt, self._to_txt)
        return fn(content, title)

    @staticmethod
    def _to_pdf(text: str, title: str) -> bytes:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import (
                Paragraph, SimpleDocTemplate, Spacer, ListFlowable,
                ListItem, Image as RLImage
            )
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

            buf = io.BytesIO()
            doc = SimpleDocTemplate(
                buf, pagesize=A4,
                rightMargin=2.5*cm, leftMargin=2.5*cm,
                topMargin=2.5*cm, bottomMargin=2.5*cm
            )

            base = getSampleStyleSheet()
            styles = {
                "doc_title": ParagraphStyle("doc_title", parent=base["Title"],
                    fontSize=22, textColor=colors.HexColor("#1F3864"),
                    spaceAfter=6, leading=28, alignment=TA_CENTER),
                "date": ParagraphStyle("date", parent=base["Normal"],
                    fontSize=9, textColor=colors.HexColor("#888888"),
                    spaceAfter=20, alignment=TA_CENTER),
                "heading": ParagraphStyle("heading", parent=base["Heading1"],
                    fontSize=14, textColor=colors.HexColor("#2E75B6"),
                    spaceBefore=16, spaceAfter=6, leading=18),
                "subheading": ParagraphStyle("subheading", parent=base["Heading2"],
                    fontSize=12, textColor=colors.HexColor("#1F3864"),
                    spaceBefore=10, spaceAfter=4, leading=16),
                "paragraph": ParagraphStyle("paragraph", parent=base["Normal"],
                    fontSize=11, textColor=colors.HexColor("#222222"),
                    spaceAfter=8, leading=17, alignment=TA_JUSTIFY),
                "bullet_item": ParagraphStyle("bullet_item", parent=base["Normal"],
                    fontSize=11, textColor=colors.HexColor("#222222"),
                    leading=16, leftIndent=0),
                "formula_label": ParagraphStyle("formula_label", parent=base["Normal"],
                    fontSize=9, textColor=colors.HexColor("#666666"),
                    alignment=TA_CENTER, spaceAfter=4),
            }

            story = [
                Paragraph(title, styles["doc_title"]),
                Paragraph(datetime.now().strftime("%d de %B de %Y"), styles["date"]),
                Spacer(1, 0.3*cm),
            ]

            blocks = _split_latex(text)
            bullet_group = []

            def flush_bullets():
                if bullet_group:
                    story.append(ListFlowable(
                        [ListItem(Paragraph(b, styles["bullet_item"]),
                                  leftIndent=20,
                                  bulletColor=colors.HexColor("#2E75B6"))
                         for b in bullet_group],
                        bulletType="bullet", leftIndent=20, spaceAfter=8
                    ))
                    bullet_group.clear()

            for block in blocks:
                t = block["type"]

                if t == "latex_block":
                    flush_bullets()
                    # Intentar renderizar como imagen
                    img_bytes = _latex_to_image(block["text"], fontsize=14, dpi=150)
                    if img_bytes:
                        img_buf = io.BytesIO(img_bytes)
                        rl_img = RLImage(img_buf, width=14*cm, height=2*cm)
                        rl_img.hAlign = "CENTER"
                        story.append(Spacer(1, 0.2*cm))
                        story.append(rl_img)
                        story.append(Spacer(1, 0.2*cm))
                    else:
                        # Fallback: mostrar como texto monoespaciado
                        story.append(Paragraph(
                            f"[Fórmula: {block['text'][:120]}]",
                            styles["formula_label"]
                        ))

                elif t == "empty":
                    flush_bullets()
                    story.append(Spacer(1, 0.2*cm))

                elif t in ("title", "heading"):
                    flush_bullets()
                    story.append(Paragraph(block["text"], styles["heading"]))

                elif t == "subheading":
                    flush_bullets()
                    story.append(Paragraph(block["text"], styles["subheading"]))

                elif t in ("bullet", "numbered"):
                    bullet_group.append(block["text"])

                elif t == "paragraph":
                    flush_bullets()
                    # Procesar inline LaTeX → imagen pequeña inline no es posible en reportlab
                    # Extraer y mostrar como texto con notación limpia
                    raw = block["text"]
                    # Convertir inline $...$ a texto entre corchetes si matplotlib no disponible
                    cleaned = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", raw)
                    cleaned = re.sub(r"\*(.*?)\*", r"<i>\1</i>", cleaned)
                    # Inline LaTeX: dejar visible pero entre espacios
                    cleaned = re.sub(r'\$([^$]+)\$', r' <i>\1</i> ', cleaned)
                    story.append(Paragraph(cleaned, styles["paragraph"]))

            flush_bullets()
            doc.build(story)
            logger.info("PDF generado: %d bytes", buf.tell())
            return buf.getvalue()

        except Exception as exc:
            logger.error("PDF generation error: %s", exc)
            return text.encode()

    @staticmethod
    def _to_docx(text: str, title: str) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import re

            doc = Document()

            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # Título
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                h.runs[0].font.size = Pt(22)

            # Fecha
            date_p = doc.add_paragraph(datetime.now().strftime("%d de %B de %Y"))
            date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if date_p.runs:
                date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                date_p.runs[0].font.size = Pt(9)
            doc.add_paragraph("")

            blocks = _split_latex(text)

            for block in blocks:
                t = block["type"]

                if t == "latex_block":
                    # Renderizar como imagen PNG embebida en el Word
                    img_bytes = _latex_to_image(block["text"], fontsize=16, dpi=200)
                    if img_bytes:
                        img_buf = io.BytesIO(img_bytes)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_buf, width=Cm(14))
                    else:
                        # Fallback: texto en cursiva centrado
                        p = doc.add_paragraph(f"[ {block['text']} ]")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p.runs:
                            p.runs[0].italic = True
                            p.runs[0].font.size = Pt(10)
                            p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x88)

                elif t == "empty":
                    doc.add_paragraph("")

                elif t in ("title", "heading"):
                    p = doc.add_heading(block["text"], level=1)
                    if p.runs:
                        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                        p.runs[0].font.size = Pt(14)

                elif t == "subheading":
                    p = doc.add_heading(block["text"], level=2)
                    if p.runs:
                        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                        p.runs[0].font.size = Pt(12)

                elif t in ("bullet", "numbered"):
                    style = "List Bullet" if t == "bullet" else "List Number"
                    p = doc.add_paragraph(block["text"], style=style)
                    if p.runs:
                        p.runs[0].font.size = Pt(11)

                elif t == "paragraph":
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    raw = block["text"]

                    # Separar inline LaTeX, negritas, itálicas
                    parts = re.split(r'(\$[^$\n]+?\$|\*\*.*?\*\*|\*.*?\*)', raw)
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith('$') and part.endswith('$') and len(part) > 2:
                            # Inline LaTeX: renderizar como imagen pequeña
                            formula = part[1:-1]
                            img_bytes = _latex_to_image(formula, fontsize=11, dpi=150)
                            if img_bytes:
                                run = p.add_run()
                                run.add_picture(io.BytesIO(img_bytes), height=Cm(0.5))
                            else:
                                run = p.add_run(f" {formula} ")
                                run.italic = True
                                run.font.size = Pt(11)
                        elif part.startswith("**") and part.endswith("**"):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                            run.font.size = Pt(11)
                        elif part.startswith("*") and part.endswith("*"):
                            run = p.add_run(part[1:-1])
                            run.italic = True
                            run.font.size = Pt(11)
                        else:
                            run = p.add_run(part)
                            run.font.size = Pt(11)

            buf = io.BytesIO()
            doc.save(buf)
            logger.info("DOCX generado")
            return buf.getvalue()

        except Exception as exc:
            logger.error("DOCX generation error: %s", exc)
            return text.encode()

    @staticmethod
    def _to_txt(text: str, title: str) -> bytes:
        header = f"{title}\n{'='*len(title)}\n{datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
        return (header + text).encode("utf-8")

    @staticmethod
    def _to_md(text: str, title: str) -> bytes:
        header = f"# {title}\n\n_{datetime.now().strftime('%Y-%m-%d %H:%M')}_\n\n"
        return (header + text).encode("utf-8")